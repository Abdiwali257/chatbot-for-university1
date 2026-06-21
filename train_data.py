"""Build / rebuild the SIMAD University chatbot knowledge base.

This script reads PDF, Word, and Excel documents, cleans them, splits them into
small searchable chunks, converts those chunks into NLP embeddings, and stores
those embeddings in ChromaDB.

This is the project's "training" step for the RAG chatbot. It does not train a
new LLM from scratch; it creates a searchable SIMAD knowledge base.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import unicodedata
from pathlib import Path

import chromadb
import pandas as pd
from dotenv import load_dotenv
from docx import Document
from pypdf import PdfReader
from sentence_transformers import SentenceTransformer

PROJECT_DIR = Path(__file__).resolve().parent
DATA_DIR = PROJECT_DIR / "data"
DB_DIR = PROJECT_DIR / "chroma_db"
COLLECTION_NAME = "simad_knowledge_base"

load_dotenv(PROJECT_DIR / ".env")

EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "all-MiniLM-L6-v2")
EMBEDDING_BATCH_SIZE = int(os.getenv("EMBEDDING_BATCH_SIZE", "16"))
CHROMA_ADD_BATCH_SIZE = int(os.getenv("CHROMA_ADD_BATCH_SIZE", "64"))
CHUNK_TARGET_SIZE = int(os.getenv("CHUNK_TARGET_SIZE", "850"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "180"))
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".json"}

# Avoid accidentally indexing thesis chapters, recovered Office lock files, or project reports.
EXCLUDED_SOURCE_NAME_PATTERN = re.compile(
    r"(?:AI_Assistant|Chapters?_4_and_5|ASP document|thesis|project report|~\$)",
    re.I,
)

SOURCE_CATEGORY_RULES = [
    ("admission", ("admission", "apply", "application", "register", "enroll", "enrol", "transfer")),
    ("fees", ("tuition", "fee", "fees", "payment", "scholarship", "financial", "qard")),
    ("facilities", ("campus", "library", "lab", "facility", "facilities", "club", "cultural", "extracurricular", "co-curricular", "disability", "dss", "exchange", "map", "service")),
    ("staff", ("rector", "senate", "dean", "staff", "lecturer", "professor", "official", "governance")),
    ("academic", ("faculty", "school", "program", "course", "curriculum", "department", "gpa", "grading", "research", "master", "postgraduate", "computing", "engineering", "medicine", "law", "economics", "education", "management", "social science")),
    ("general", ("history", "vision", "mission", "why simad", "general information", "accreditation", "ranking", "membership", "xajsi")),
]


def clean_text(value: object) -> str:
    text = unicodedata.normalize("NFKC", str(value or ""))
    text = text.replace("\u200b", " ").replace("\ufb01", "fi").replace("\ufb02", "fl")
    return " ".join(text.split())


def normalize_document_text(value: object) -> str:
    text = unicodedata.normalize("NFKC", str(value or ""))
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\ufb01", "fi").replace("\ufb02", "fl")
    lines = []
    for raw_line in text.split("\n"):
        line = clean_text(raw_line)
        if line:
            lines.append(line)
    return "\n".join(lines).strip()


def classify_chunk_category(path: Path, section: str, text: str) -> str:
    haystack = f"{path.name} {section} {text}".lower()
    for category, keywords in SOURCE_CATEGORY_RULES:
        if any(keyword in haystack for keyword in keywords):
            return category
    return "general"


def document_files() -> list[Path]:
    """Return source documents from data/ and useful root-level faculty files."""
    files: list[Path] = []
    if DATA_DIR.exists():
        files.extend(DATA_DIR.rglob("*"))
    files.extend(PROJECT_DIR.glob("*.docx"))
    files.extend(PROJECT_DIR.glob("*.xlsx"))
    files.extend(PROJECT_DIR.glob("*.xls"))
    files.extend(PROJECT_DIR.glob("simad_official_public_info.json"))
    result = []
    for path in files:
        if not path.is_file():
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        if EXCLUDED_SOURCE_NAME_PATTERN.search(path.name):
            continue
        result.append(path.resolve())
    return sorted(set(result))


def read_pdf(path: Path) -> list[tuple[str, dict[str, object]]]:
    records = []
    reader = PdfReader(str(path))
    for page_number, page in enumerate(reader.pages, start=1):
        text = normalize_document_text(page.extract_text() or "")
        if text:
            records.append((text, {"location": f"page {page_number}", "page_number": page_number}))
    return records


def read_docx(path: Path) -> list[tuple[str, dict[str, object]]]:
    doc = Document(str(path))
    blocks = []
    for para in doc.paragraphs:
        text = normalize_document_text(para.text)
        if text:
            blocks.append(text)
    for table_number, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells if clean_text(cell.text)]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            blocks.append(f"Table {table_number}:\n" + "\n".join(rows))
    text = "\n".join(blocks)
    return [(text, {"location": "document"})] if text else []


def read_excel(path: Path) -> list[tuple[str, dict[str, object]]]:
    records = []
    sheets = pd.read_excel(path, sheet_name=None, header=None)
    for sheet_name, frame in sheets.items():
        rows = []
        for values in frame.fillna("").itertuples(index=False, name=None):
            cells = [clean_text(value) for value in values if clean_text(value)]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            records.append(("\n".join(rows), {"location": f"sheet {sheet_name}", "section": str(sheet_name)}))
    return records


def read_json(path: Path) -> list[tuple[str, dict[str, object]]]:
    if path.name != "simad_official_public_info.json":
        return []
    data = json.loads(path.read_text(encoding="utf-8"))
    records = []
    for school in data.get("schools", []):
        name = clean_text(school.get("name", ""))
        lines = [f"Official SIMAD school: {name}"]
        if school.get("email"):
            lines.append(f"Email: {clean_text(school['email'])}")
        if school.get("website"):
            lines.append(f"Website: {clean_text(school['website'])}")
        programs = [clean_text(program) for program in school.get("programs", []) if clean_text(program)]
        if programs:
            lines.append("Programs: " + ", ".join(programs))

        administration = school.get("administration") or []
        if administration:
            lines.append("Administration:")
            for member in administration:
                detail = f"- {clean_text(member.get('name', ''))}: {clean_text(member.get('role', ''))}"
                if member.get("email"):
                    detail += f"; email: {clean_text(member['email'])}"
                if member.get("phone"):
                    detail += f"; phone: {clean_text(member['phone'])}"
                lines.append(detail)
        elif school.get("administration_status"):
            lines.append(f"Administration status: {clean_text(school['administration_status'])}")

        source = clean_text(school.get("administration_source", "official SIMAD public information JSON"))
        records.append((
            "\n".join(lines),
            {"location": source, "section": name or "Official SIMAD public information"},
        ))

    for item in data.get("secondary_or_possibly_stale_mentions", []):
        mention = clean_text(item.get("mention", ""))
        if mention:
            records.append((
                f"Secondary SIMAD staff mention needing verification: {mention}. "
                f"Note: {clean_text(item.get('note', ''))}",
                {
                    "location": clean_text(item.get("source", "secondary official source")),
                    "section": "Secondary staff mentions",
                },
            ))
    return records


def read_records(path: Path) -> list[tuple[str, dict[str, object]]]:
    if path.suffix.lower() == ".pdf":
        return read_pdf(path)
    if path.suffix.lower() == ".docx":
        return read_docx(path)
    if path.suffix.lower() in {".xlsx", ".xls"}:
        return read_excel(path)
    if path.suffix.lower() == ".json":
        return read_json(path)
    return []


def source_name(path: Path) -> str:
    try:
        return str(path.relative_to(PROJECT_DIR))
    except ValueError:
        return path.name


def split_long_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    compact = clean_text(text)
    if not compact:
        return []
    if len(compact) <= chunk_size:
        return [compact]
    chunks = []
    start = 0
    while start < len(compact):
        end = min(start + chunk_size, len(compact))
        if end < len(compact):
            lower_bound = start + int(chunk_size * 0.60)
            break_at = max(
                compact.rfind(". ", lower_bound, end),
                compact.rfind("; ", lower_bound, end),
                compact.rfind(" | ", lower_bound, end),
                compact.rfind(" ", lower_bound, end),
            )
            if break_at > lower_bound:
                end = break_at + 1
        chunk = compact[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(compact):
            break
        start = max(end - overlap, start + 1)
    return chunks


def build_knowledge_base() -> None:
    print(f"Loading NLP embedding model: {EMBEDDING_MODEL}", flush=True)
    model = SentenceTransformer(EMBEDDING_MODEL)
    client = chromadb.PersistentClient(path=str(DB_DIR))
    try:
        client.delete_collection(COLLECTION_NAME)
    except Exception:
        pass
    collection = client.create_collection(
        COLLECTION_NAME,
        metadata={"hnsw:space": "cosine", "embedding_model": EMBEDDING_MODEL},
    )

    documents: list[str] = []
    metadatas: list[dict[str, object]] = []
    ids: list[str] = []
    loaded_files = 0

    for path in document_files():
        try:
            records = read_records(path)
            file_chunks = 0
            for text, record_metadata in records:
                for chunk_number, chunk in enumerate(split_long_text(text, CHUNK_TARGET_SIZE, CHUNK_OVERLAP), start=1):
                    section = str(record_metadata.get("section", "General"))
                    metadata = {
                        "source": source_name(path),
                        "source_file": path.name,
                        "file_type": path.suffix.lower().lstrip("."),
                        "location": str(record_metadata.get("location", "document")),
                        "section": section,
                        "page_number": int(record_metadata.get("page_number", 0) or 0),
                        "chunk": chunk_number,
                        "category": classify_chunk_category(path, section, chunk),
                    }
                    digest = hashlib.sha1(
                        f"{metadata['source']}|{metadata['location']}|{section}|{chunk_number}|{chunk}".encode("utf-8")
                    ).hexdigest()
                    documents.append(chunk)
                    metadatas.append(metadata)
                    ids.append(digest)
                    file_chunks += 1
            if file_chunks:
                loaded_files += 1
                print(f"Loaded: {source_name(path)} | Chunks: {file_chunks}", flush=True)
            else:
                print(f"Skipped empty file: {source_name(path)}", flush=True)
        except Exception as exc:
            print(f"Error reading {source_name(path)}: {exc}", flush=True)

    if not documents:
        raise SystemExit("No supported documents with extractable text were found.")

    for start in range(0, len(documents), CHROMA_ADD_BATCH_SIZE):
        end = min(start + CHROMA_ADD_BATCH_SIZE, len(documents))
        embeddings = model.encode(
            documents[start:end],
            batch_size=EMBEDDING_BATCH_SIZE,
            normalize_embeddings=True,
            show_progress_bar=False,
        ).tolist()
        collection.add(
            documents=documents[start:end],
            embeddings=embeddings,
            ids=ids[start:end],
            metadatas=metadatas[start:end],
        )
        print(f"Indexed {end}/{len(documents)} chunks", flush=True)

    print(f"\nKnowledge base ready: {loaded_files} files, {len(documents)} chunks.", flush=True)
    print("Training completed successfully.", flush=True)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__)
    parser.parse_args()
    build_knowledge_base()
