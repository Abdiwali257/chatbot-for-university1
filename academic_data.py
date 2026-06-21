"""Structured SIMAD academic and tuition data extracted from source documents."""

from __future__ import annotations

import json
import re
from dataclasses import dataclass, replace
from difflib import get_close_matches
from functools import lru_cache
from pathlib import Path

import pandas as pd
from docx import Document
from pypdf import PdfReader

PROJECT_DIR = Path(__file__).resolve().parent
SEMESTERS = {
    "one": 1,
    "first": 1,
    "two": 2,
    "second": 2,
    "three": 3,
    "third": 3,
    "four": 4,
    "fourth": 4,
    "five": 5,
    "fifth": 5,
    "six": 6,
    "sixth": 6,
    "seven": 7,
    "seventh": 7,
    "eight": 8,
    "eighth": 8,
    "nine": 9,
    "ninth": 9,
    "ten": 10,
    "tenth": 10,
}


def normalize(value: object) -> str:
    return " ".join(str(value).replace("\t", " ").split())


def semester_number(text: str) -> int | None:
    words = "|".join(SEMESTERS)
    match = re.search(rf"semester\s+({words}|\d+)", text, re.I)
    if not match:
        match = re.search(rf"\b({words})\s+semester\b", text, re.I)
    if not match:
        return None
    value = match.group(1).lower()
    return int(value) if value.isdigit() else SEMESTERS[value]


def looks_like_code(value: str) -> bool:
    compact = value.replace(" ", "")
    return bool(re.fullmatch(r"[A-Za-z]{2,8}[-]?\d{2,5}[A-Za-z]?", compact))


@dataclass(frozen=True)
class Course:
    faculty: str
    program: str
    semester: int
    code: str
    title: str
    credits: str
    theory: str
    practice: str


@dataclass(frozen=True)
class Tuition:
    program: str
    years: int
    fee: int
    charges: int
    total: int


@dataclass(frozen=True)
class AdministrationMember:
    school: str
    name: str
    role: str
    email: str = ""
    phone: str = ""
    source: str = ""
    note: str = ""


@dataclass(frozen=True)
class SchoolOfficialInfo:
    name: str
    email: str
    website: str
    slug: str
    programs: tuple[str, ...]
    administration: tuple[AdministrationMember, ...]
    administration_status: str = ""
    administration_source: str = ""


def courses_from_xlsx(path: Path) -> list[Course]:
    courses = []
    faculty = path.stem.strip()
    for sheet_name, frame in pd.read_excel(path, sheet_name=None, header=None).items():
        semester = 0
        for values in frame.fillna("").itertuples(index=False, name=None):
            cells = [normalize(value) for value in values]
            row_text = " ".join(cells)
            semester = semester_number(row_text) or semester

            code_index = next((i for i, value in enumerate(cells) if looks_like_code(value)), None)
            if code_index is None or code_index + 2 >= len(cells):
                continue
            title = cells[code_index + 1]
            credits = cells[code_index + 2]
            if not title or not credits:
                continue
            courses.append(
                Course(
                    faculty=faculty,
                    program=normalize(sheet_name),
                    semester=semester,
                    code=cells[code_index],
                    title=title,
                    credits=credits,
                    theory=cells[code_index + 3] if code_index + 3 < len(cells) else "",
                    practice=cells[code_index + 4] if code_index + 4 < len(cells) else "",
                )
            )
    return courses


def courses_from_docx(path: Path) -> list[Course]:
    document = Document(str(path))
    paragraphs = [normalize(p.text) for p in document.paragraphs if normalize(p.text)]
    programs = [
        text
        for text in paragraphs
        if not text.lower().startswith(("semester ", "faculty ", "undergraduate courses"))
    ]
    if not programs:
        programs = [path.stem.strip()]

    courses = []
    tables_per_program = max(1, len(document.tables) // len(programs))
    for table_index, table in enumerate(document.tables):
        program_index = min(table_index // tables_per_program, len(programs) - 1)
        semester = (table_index % tables_per_program) + 1
        for row in table.rows:
            cells = [normalize(cell.text) for cell in row.cells]
            if len(cells) < 3 or not looks_like_code(cells[0]):
                continue
            courses.append(
                Course(
                    faculty=path.stem.strip(),
                    program=programs[program_index],
                    semester=semester,
                    code=cells[0],
                    title=cells[1],
                    credits=cells[2],
                    theory=cells[3] if len(cells) > 3 else "",
                    practice=cells[4] if len(cells) > 4 else "",
                )
            )
    return courses


@lru_cache(maxsize=1)
def all_courses() -> list[Course]:
    courses = []
    for path in sorted(PROJECT_DIR.glob("*.xlsx")):
        if path.name.startswith("~$"):
            continue
        courses.extend(courses_from_xlsx(path))
    for path in sorted(PROJECT_DIR.glob("*.docx")):
        if path.name.startswith("~$"):
            continue
        courses.extend(courses_from_docx(path))
    return courses


@lru_cache(maxsize=1)
def tuition_records() -> list[Tuition]:
    text = "\n".join(
        page.extract_text() or "" for page in PdfReader(str(PROJECT_DIR / "data" / "TUTION FEES.pdf")).pages
    )
    records = []
    pattern = re.compile(r"(Bachelor of .+?)\s+(\d+)\s+(\d+)\s+(\d+)\s+(\d+)(?=\s+Bachelor|\s+School|\s+Home|\s+ENGLISH)", re.I)
    for program, years, fee, charges, total in pattern.findall(text):
        records.append(
            Tuition(normalize(program), int(years), int(fee), int(charges), int(total))
        )
    return records


@lru_cache(maxsize=1)
def faculty_programs() -> dict[str, tuple[str, ...]]:
    """Return verified degree programs grouped by school/faculty from the tuition document."""
    text = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(str(PROJECT_DIR / "data" / "TUTION FEES.pdf")).pages
    )
    corrections = {
        "Computer Scienece": "Computer Science",
        "Econimics": "Economics",
        "Enterpreneurship": "Entrepreneurship",
    }
    grouped = {}
    section_pattern = re.compile(
        r"School of ([^\r\n]+)\s+(.*?)(?=School of |ENGLISH SKILLS PROGRAM|Academic Programs)",
        re.I | re.S,
    )
    program_pattern = re.compile(r"Bachelor of (.+?)\s+\d+\s+\d+\s+\d+\s+\d+", re.I)
    for faculty, section in section_pattern.findall(text):
        label = normalize(faculty)
        programs = []
        for program in program_pattern.findall(section):
            cleaned = normalize(program)
            for wrong, corrected in corrections.items():
                cleaned = cleaned.replace(wrong, corrected)
            programs.append(cleaned)
        if programs:
            grouped[label] = tuple(programs)
    return grouped


def find_programs(question: str) -> tuple[str, tuple[str, ...]] | None:
    """Find a verified faculty and its degree programs without using course/module rows."""
    normalized_question = normalize_academic_query(question)
    wanted = query_terms(normalized_question) - {
        "faculty",
        "school",
        "of",
        "program",
        "programs",
        "degree",
        "degrees",
    }
    scored = []
    for faculty, programs in faculty_programs().items():
        overlap = len(wanted & query_terms(faculty))
        if overlap:
            scored.append((overlap, faculty, programs))
    if not scored:
        return None
    _, faculty, programs = max(scored, key=lambda item: item[0])
    return faculty, programs


def find_program_parent(question: str) -> tuple[str, str] | None:
    """Return the verified faculty containing a named undergraduate program."""
    normalized_question = normalize_academic_query(question)
    lowered = normalized_question.lower()
    aliases = {
        "cs": "Computer Science",
        "it": "Information Technology",
        "gm": "Graphics and Multimedia",
    }
    for alias, program in aliases.items():
        if re.search(rf"\b{alias}\b", lowered):
            for faculty, programs in faculty_programs().items():
                if program in programs:
                    return faculty, program

    wanted = query_terms(normalized_question)
    candidates = []
    for faculty, programs in faculty_programs().items():
        for program in programs:
            program_terms = query_terms(program)
            if program_terms and program_terms <= wanted:
                candidates.append((len(program_terms), faculty, program))
            else:
                overlap = wanted & program_terms
                if len(overlap) >= 2 and len(overlap) / max(1, len(program_terms)) >= 0.5:
                    candidates.append((len(overlap), faculty, program))
    if not candidates:
        return None
    _, faculty, program = max(candidates, key=lambda item: item[0])
    return faculty, program


def query_terms(text: str) -> set[str]:
    return set(re.findall(r"[a-z0-9]+", text.lower()))


QUERY_CORRECTIONS = {
    "acreditation": "accreditation",
    "acredition": "accreditation",
    "adivce": "advice",
    "admissom": "admission",
    "admisson": "admission",
    "admisions": "admissions",
    "aplication": "application",
    "bacheler": "bachelor",
    "carrer": "career",
    "clbus": "clubs",
    "computr": "computer",
    "disabilty": "disability",
    "econimics": "economics",
    "enginering": "engineering",
    "faclty": "faculty",
    "facultly": "faculty",
    "infomation": "information",
    "inteligence": "intelligence",
    "libary": "library",
    "medecine": "medicine",
    "programms": "programs",
    "progrms": "programs",
    "recomend": "recommend",
    "registeration": "registration",
    "requirments": "requirements",
    "scholrship": "scholarship",
    "scholrships": "scholarships",
    "scince": "science",
    "semster": "semester",
    "simd": "simad",
    "subjcts": "subjects",
    "tution": "tuition",
    "universty": "university",
}
DOMAIN_QUERY_WORDS = {
    "accommodation",
    "accreditation",
    "academic",
    "advice",
    "admission",
    "admissions",
    "application",
    "apply",
    "bachelor",
    "campus",
    "career",
    "choose",
    "clubs",
    "conference",
    "conferences",
    "conduct",
    "course",
    "courses",
    "credit",
    "credits",
    "curriculum",
    "cultural",
    "degree",
    "degrees",
    "department",
    "departments",
    "disability",
    "discipline",
    "disclosure",
    "exchange",
    "extracurricular",
    "faculty",
    "faculties",
    "fees",
    "grading",
    "history",
    "hours",
    "innovation",
    "institute",
    "languages",
    "library",
    "membership",
    "memberships",
    "mission",
    "module",
    "modules",
    "program",
    "programs",
    "postgraduate",
    "ranking",
    "rankings",
    "recommend",
    "rector",
    "register",
    "registration",
    "requirements",
    "research",
    "scholarship",
    "scholarships",
    "school",
    "senate",
    "semester",
    "services",
    "simad",
    "student",
    "students",
    "subject",
    "subjects",
    "tuition",
    "university",
    "values",
    "vision",
}
PROTECTED_QUERY_WORDS = {
    "about",
    "also",
    "and",
    "are",
    "can",
    "could",
    "does",
    "explain",
    "for",
    "from",
    "give",
    "have",
    "help",
    "inside",
    "into",
    "list",
    "need",
    "offer",
    "offers",
    "please",
    "show",
    "tell",
    "that",
    "their",
    "there",
    "these",
    "they",
    "this",
    "those",
    "want",
    "what",
    "when",
    "where",
    "which",
    "who",
    "with",
    "would",
}


@lru_cache(maxsize=1)
def academic_vocabulary() -> tuple[str, ...]:
    """Build a conservative spelling vocabulary from verified academic records."""
    vocabulary = set(DOMAIN_QUERY_WORDS)
    vocabulary.update(SEMESTERS)
    for faculty, programs in faculty_programs().items():
        vocabulary.update(query_terms(faculty))
        for program in programs:
            vocabulary.update(query_terms(program))
    for course in all_courses():
        vocabulary.update(query_terms(course.faculty))
        vocabulary.update(query_terms(course.program))
        vocabulary.update(query_terms(course.title))
    return tuple(sorted(word for word in vocabulary if len(word) >= 4))


def normalize_academic_query(text: str) -> str:
    """Correct likely SIMAD-domain misspellings while leaving names and codes alone."""
    text = re.sub(r"\bfoc\b", "Faculty of Computing", text, flags=re.I)
    vocabulary = academic_vocabulary()

    def replace_word(match: re.Match[str]) -> str:
        original = match.group(0)
        lowered = original.lower()
        if (
            len(lowered) < 4
            or lowered in PROTECTED_QUERY_WORDS
            or lowered in vocabulary
            or (original.isupper() and len(original) <= 8)
        ):
            return original

        corrected = QUERY_CORRECTIONS.get(lowered)
        if not corrected:
            candidates = [
                word
                for word in vocabulary
                if word[0] == lowered[0] and abs(len(word) - len(lowered)) <= 2
            ]
            matches = get_close_matches(lowered, candidates, n=1, cutoff=0.86)
            corrected = matches[0] if matches else None
        if not corrected:
            return original
        return corrected.capitalize() if original[0].isupper() else corrected

    return re.sub(r"[A-Za-z]+", replace_word, text)


COMPUTING_PROGRAM_MARKERS = {
    "Computer Science": "CS",
    "Information Technology": "IT",
    "Graphics and Multimedia": "GM",
}
COMPUTING_PROGRAM_ALIASES = {
    "Computer Science": ("computer science", "cs"),
    "Information Technology": ("information technology", "it"),
    "Graphics and Multimedia": ("graphics and multimedia", "graphic and multimedia", "gm"),
}
GENERIC_PROGRAM_TERMS = {
    "department",
    "faculty",
    "program",
    "school",
    "of",
}


def display_program_name(value: str) -> str:
    cleaned = re.sub(r"^Department of\s+", "", normalize(value), flags=re.I)
    cleaned = re.sub(r"\s+(?:department|program)$", "", cleaned, flags=re.I)
    cleaned = re.sub(r"\bMangem?ents\b", "Management", cleaned, flags=re.I)
    return cleaned.title() if cleaned.isupper() else cleaned


SCHOOL_ALIAS_TERMS = {
    "School of Accountancy": ("fa", "accountancy", "accounting"),
    "School of Computing": ("foc", "fc", "computing", "computer science"),
    "School of Economics": ("feco", "economics", "statistics"),
    "School of Education": ("fedu", "education"),
    "School of Engineering": ("feng", "engineering", "civil engineering", "architecture"),
    "School of Law": ("flaw", "law"),
    "School of Management Sciences": ("fms", "management", "business"),
    "School of Medicine and Health Sciences": (
        "fmhs",
        "medicine",
        "health sciences",
        "nursing",
        "public health",
    ),
    "School of Social Sciences": ("fass", "social sciences", "political science"),
    "SIMAD Master Programs": ("graduate school", "gs", "master", "postgraduate"),
    "OUM Programs": ("oum",),
}


def compact_staff_role(role: str) -> str:
    return normalize(role).replace("HOD", "Head of Department")


@lru_cache(maxsize=1)
def simad_official_schools() -> tuple[SchoolOfficialInfo, ...]:
    path = PROJECT_DIR / "simad_official_public_info.json"
    if not path.exists():
        return ()
    data = json.loads(path.read_text(encoding="utf-8"))
    schools = []
    for item in data.get("schools", []):
        school_name = normalize(item.get("name", ""))
        source = normalize(item.get("administration_source", ""))
        administration = tuple(
            AdministrationMember(
                school=school_name,
                name=normalize(member.get("name", "")),
                role=compact_staff_role(member.get("role", "")),
                email=normalize(member.get("email", "")),
                phone=normalize(member.get("phone", "")),
                source=source,
                note=normalize(member.get("note", "")),
            )
            for member in item.get("administration", []) or []
            if normalize(member.get("name", "")) and normalize(member.get("role", ""))
        )
        schools.append(
            SchoolOfficialInfo(
                name=school_name,
                email=normalize(item.get("email", "")),
                website=normalize(item.get("website", "")),
                slug=normalize(item.get("slug", "")),
                programs=tuple(
                    normalize(program)
                    for program in item.get("programs", []) or []
                    if normalize(program)
                ),
                administration=administration,
                administration_status=normalize(item.get("administration_status", "")),
                administration_source=source,
            )
        )
    return tuple(schools)


def official_school_labels(school: SchoolOfficialInfo) -> tuple[str, ...]:
    base = school.name
    short = re.sub(r"^School of\s+", "", base, flags=re.I)
    labels = [base, short, f"Faculty of {short}", school.slug.replace("-", " ")]
    labels.extend(SCHOOL_ALIAS_TERMS.get(base, ()))
    labels.extend(school.programs)
    labels.extend(member.role for member in school.administration)
    return tuple(label for label in labels if normalize(label))


def find_official_school(question: str) -> SchoolOfficialInfo | None:
    normalized_question = normalize_academic_query(question)
    lowered = normalized_question.lower()
    wanted = query_terms(normalized_question) - {
        "simad", "university", "officials", "official", "people",
        "charge", "top", "leader", "leaders", "leadership", "who", "are",
        "the", "in", "of", "runs", "manages", "institution", "and", "or",
        "with", "for", "to", "at", "on", "by", "an", "a", "any", "some"
    }
    candidates = []
    for school in simad_official_schools():
        score = 0
        for label in official_school_labels(school):
            label_terms = query_terms(label) - GENERIC_PROGRAM_TERMS
            if not label_terms:
                continue
            label_text = label.lower()
            if re.search(rf"\b{re.escape(label_text)}\b", lowered):
                score = max(score, 20 + len(label_terms))
            elif label_terms <= wanted:
                score = max(score, 12 + len(label_terms))
            else:
                overlap = wanted & label_terms
                if overlap:
                    score = max(score, len(overlap))
        if score:
            candidates.append((score, school.name, school))
    if not candidates:
        return None
    return max(candidates, key=lambda item: (item[0], item[1]))[2]


def member_matches_query(member: AdministrationMember, lowered_question: str) -> bool:
    role = member.role.lower()
    if re.search(r"\bassistant\s+dean\b", lowered_question):
        return "assistant dean" in role
    if re.search(r"\bdeans?\b", lowered_question):
        return role.startswith("dean") or re.match(r"^dean\b", role)
    if re.search(r"\b(?:head|heads|hod)\b", lowered_question):
        return "head" in role or "hod" in role
    if "coordinator" in lowered_question:
        return "coordinator" in role
    if "secretary" in lowered_question:
        return "secretary" in role
    if "lecturer" in lowered_question:
        return "lecturer" in role
    return True


def administration_context(question: str) -> str:
    lowered = normalize_academic_query(question).lower()
    asks_admin = re.search(
        r"\b(?:administration|administrator|deans?|heads?|hod|staff|lecturers?|"
        r"coordinators?|secretary|in charge|officials?|leadership)\b",
        lowered,
    )
    if not asks_admin:
        return ""

    if re.search(
        r"\b(?:all|list|show)\b.*\bdeans?\b|\bdeans\b.*\bfacult|\bfaculty\s+deans\b",
        lowered,
    ):
        lines = ["Verified SIMAD administration records:"]
        for school in simad_official_schools():
            deans = [
                member
                for member in school.administration
                if member.role.lower().startswith("dean")
            ]
            if deans:
                for dean in deans:
                    lines.append(f"- {school.name}: {dean.name} - {dean.role}")
            elif school.administration_status:
                lines.append(f"- {school.name}: {school.administration_status}")
        return "\n".join(lines) if len(lines) > 1 else ""

    school = find_official_school(question)
    if not school:
        return ""

    members = [
        member
        for member in school.administration
        if member_matches_query(member, lowered)
    ]
    if members and re.search(r"\b(?:head|heads|hod)\b", lowered):
        generic = {
            "department",
            "faculty",
            "head",
            "heads",
            "hod",
            "of",
            "school",
            "who",
        }
        wanted = query_terms(lowered) - generic
        specific_members = [
            member
            for member in members
            if (query_terms(member.role) - generic) & wanted
        ]
        if specific_members:
            members = specific_members
    if not members and school.administration:
        members = list(school.administration)

    lines = ["Verified SIMAD administration records:", f"School: {school.name}"]
    if members:
        for member in members:
            detail = f"- {member.name} - {member.role}"
            contacts = []
            if member.email:
                contacts.append(f"email: {member.email}")
            if member.phone:
                contacts.append(f"phone: {member.phone}")
            if contacts:
                detail += f" ({'; '.join(contacts)})"
            lines.append(detail)
    elif school.administration_status:
        lines.append(school.administration_status)
    else:
        lines.append("No current administration records were found for this school.")
    return "\n".join(lines)


def course_program_name(question: str) -> str | None:
    normalized_question = normalize_academic_query(question)
    lowered = normalized_question.lower()
    for program, aliases in COMPUTING_PROGRAM_ALIASES.items():
        if any(re.search(rf"\b{re.escape(alias)}\b", lowered) for alias in aliases):
            return program

    wanted = query_terms(normalized_question)
    candidates = []
    for raw_program in {course.program for course in all_courses()}:
        display = display_program_name(raw_program)
        label_terms = query_terms(display) - GENERIC_PROGRAM_TERMS
        if label_terms and label_terms <= wanted:
            candidates.append((len(label_terms), display))
    return max(candidates, default=(0, None), key=lambda item: item[0])[1]


def clean_course_title(title: str) -> str:
    cleaned = normalize(
        re.sub(
            r"\s*\(\s*(?:IT|CS|GM)(?:\s*/\s*(?:IT|CS|GM))*\s*\)\s*",
            " ",
            title,
            flags=re.I,
        )
    )
    corrections = {
        "Maintanace": "Maintenance",
        "Infromation": "Information",
        "Stucture": "Structure",
        "Architechture": "Architecture",
    }
    for wrong, corrected in corrections.items():
        cleaned = cleaned.replace(wrong, corrected)
    return cleaned


def computing_course_applies(course: Course, program: str) -> bool:
    markers = set(re.findall(r"\b(?:IT|CS|GM)\b", course.title.upper()))
    return not markers or COMPUTING_PROGRAM_MARKERS[program] in markers


def find_courses(question: str) -> list[Course]:
    normalized_question = normalize_academic_query(question)
    target_program = course_program_name(normalized_question)
    if not target_program:
        return []

    semester = semester_number(normalized_question)
    matches = []
    for course in all_courses():
        if semester is not None and course.semester != semester:
            continue

        if target_program in COMPUTING_PROGRAM_MARKERS:
            if course.faculty != "Faculty of Computing":
                continue
            if not computing_course_applies(course, target_program):
                continue
        elif display_program_name(course.program).lower() != target_program.lower():
            continue

        matches.append(
            replace(course, program=target_program, title=clean_course_title(course.title))
        )
    return matches


def find_named_courses(question: str) -> list[Course]:
    """Find course records explicitly named in a question."""
    wanted = query_terms(normalize_academic_query(question)) - {
        "about",
        "course",
        "courses",
        "i",
        "information",
        "mean",
        "module",
        "subject",
        "tell",
        "the",
    }
    if not wanted:
        return []
    scored = []
    for course in all_courses():
        title = clean_course_title(course.title)
        title_terms = query_terms(title)
        overlap = len(wanted & title_terms)
        if overlap >= 2 and overlap / len(wanted) >= 0.5:
            scored.append((overlap / len(title_terms), replace(course, title=title)))
    if not scored:
        return []
    best = max(score for score, _ in scored)
    return [course for score, course in scored if score == best][:8]
